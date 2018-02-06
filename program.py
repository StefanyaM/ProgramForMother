# -*- coding: utf-8 -*- 

import re
import string
from collections import Counter
from docx import Document

# patterns
punc = r'[{}]?$'.format(string.punctuation)
chik_pattern = r'^\w{3,}(чик|щик)\w{,3}' + punc
ets_pattern = r'^\w{3,}(ец|иц|[^аоуэияюы]ц)\w{,3}' + punc
ik_pattern = r'^\w{3,}([^чщ]ик|ник)\w{,3}' + punc
un_pattern = r'^\w{3,}(ун)\w{,3}' + punc
yak_pattern = r'^\w{3,}(як|ак)\w{,3}' + punc
an_pattern = r'^\w{3,}ан([аминовуы]{,3})' + punc
ist_pattern = r'^\w{2,}ист\w{,3}' + punc
tel_pattern = r'^\w{3,}тел\w{1,3}' + punc
ant_pattern = r'^\w{4,}(ант|ент)\w{,3}' + punc
or_pattern = r'^\w{3,}(ор|ер|ар|яр|ир)([аминовуы]{,3})' + punc
uh_pattern = r'^\w{3,}ух\w{,1}' + punc
al_pattern = r'^\w{3,}ал([аминовуы]{,3})' + punc

patterns = {'чик/щик': chik_pattern, 'ец/иц': ets_pattern, 'ик/ник' : ik_pattern, 'ун' : un_pattern, 'як/ак' : yak_pattern, 'ан' : an_pattern, 'ист' : ist_pattern, 'тель' : tel_pattern, 'ант/ент' : ant_pattern, 'ор/ер/ар/яр/ир' : or_pattern, 'ух' : uh_pattern, 'ал' : al_pattern}

# program
document = Document ()
p = document.add_paragraph()
words = {key: [] for key in patterns.keys()}
    
file = open ('text.txt', 'r', encoding ='utf-8')
f = file.read().split()
file.close()
    
for word in f:
    bold_index = 0
        
    for name, pat in patterns.items():
        if bool (re.match (pat, word, re.IGNORECASE)) == True:
            words [name].append (re.search (r'\w+', word).group().lower())
            bold_index = 1
                
	if bold_index == 0:
		p.add_run (word + ' ')
    else:
		p.add_run (word + ' ').bold = True
    
document.add_paragraph (' ')
for key, val in words.items():
	document.add_paragraph (key, style ='ListNumber')
	for elem in Counter (val).most_common():
		document.add_paragraph (elem[0] + ' ' + str(elem[1]))
            
document.save (i[:-4] + '.docx')
        
print ('All done')    
